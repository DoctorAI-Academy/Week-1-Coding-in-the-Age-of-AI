# Create a Word template that mirrors the user's provided structure.
# If python-docx isn't available, create an RTF fallback with the same headings.
from datetime import date
from pathlib import Path

docx_path = Path("/mnt/data/Proposed_Plan_of_Study_Template.docx")
rtf_path = Path("/mnt/data/Proposed_Plan_of_Study_Template.rtf")

def make_docx(path):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # --- Title ---
    title = doc.add_paragraph()
    run = title.add_run("PROPOSED PLAN OF STUDY")
    run.bold = True
    run.font.size = Pt(18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"Generated: {date.today().isoformat()}")
    sub_run.italic = True
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    def add_section(header_text, foot_total_placeholder):
        # Header
        p = doc.add_paragraph()
        r = p.add_run(header_text)
        r.bold = True
        r.font.size = Pt(12)

        # Table
        t = doc.add_table(rows=1, cols=5)
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        t.style = "Light Grid"
        hdr = t.rows[0].cells
        hdr[0].text = "COURSE"
        hdr[1].text = "TITLE"
        hdr[2].text = "CREDITS"
        hdr[3].text = "GRADE"
        hdr[4].text = "DATE"

        # Example/placeholder row(s)
        for _ in range(3):
            row = t.add_row().cells
            row[0].text = "{{CODE}}"
            row[1].text = "{{TITLE}}"
            row[2].text = "{{CREDITS}}"
            row[3].text = "{{GRADE}}"
            row[4].text = "{{TERM YEAR}}"

        # Total line
        doc.add_paragraph("")
        pr = doc.add_paragraph()
        pr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rr = pr.add_run(f"Total\t\t{foot_total_placeholder}")
        rr.bold = True
        doc.add_paragraph("")

    # Sections A-D as in the user's structure
    add_section("A.  Core ({{CORE_TOTAL}} credits)", "{{CORE_TOTAL}}")
    add_section("B.  Additional Required Core for Measurement and Quantitative Methods Concentration ({{ARC_TOTAL}} Hours)", "{{ARC_TOTAL}}")
    add_section("C.  Electives ({{ELECTIVE_TOTAL}} Hours)", "{{ELECTIVE_TOTAL}}")
    add_section("D.  Dissertation Hours ({{DISS_TOTAL}})", "{{DISS_TOTAL}}")

    # --- Summary of Proposed Course of Study Credit Hours ---
    doc.add_paragraph("")
    s_hdr = doc.add_paragraph()
    s_run = s_hdr.add_run("SUMMARY OF PROPOSED COURSE OF STUDY CREDIT HOURS")
    s_run.bold = True

    st = doc.add_table(rows=5, cols=2)
    st.style = "Light Grid"
    st.cell(0,0).text = "Core"
    st.cell(0,1).text = "{{CORE_TOTAL}}"
    st.cell(1,0).text = "Concentration Core"
    st.cell(1,1).text = "{{ARC_TOTAL}}"
    st.cell(2,0).text = "Electives"
    st.cell(2,1).text = "{{ELECTIVE_TOTAL}}"
    st.cell(3,0).text = "Dissertation"
    st.cell(3,1).text = "{{DISS_TOTAL}}"
    st.cell(4,0).text = "Total"
    st.cell(4,1).text = "{{OVERALL_TOTAL}}"

    # --- Projected Timetable ---
    doc.add_paragraph("")
    t_hdr = doc.add_paragraph()
    t_run = t_hdr.add_run("PROJECTED TIMETABLE")
    t_run.bold = True

    tt = doc.add_table(rows=8, cols=2)
    tt.style = "Light Grid"
    rows = [
        ("Date Admitted:", "{{TERM YEAR}}"),
        ("Hours completed to date", "{{HOURS_COMPLETED}}"),
        ("Completion of Coursework", "{{TERM YEAR}}"),
        ("Comprehensive exam", "{{TERM YEAR}}"),
        ("Dissertation proposal", "{{TERM YEAR}}"),
        ("Semesters of Residency", "{{RESIDENCY_SEMESTERS}}"),
        ("Completion of dissertation", "{{TERM YEAR}}"),
        ("Notes", "{{NOTES}}"),
    ]
    for i, (k, v) in enumerate(rows):
        tt.cell(i,0).text = k
        tt.cell(i,1).text = v

    # Footer
    doc.add_paragraph("")
    note = doc.add_paragraph()
    nrun = note.add_run("Placeholders appear as {{LIKE_THIS}}. Replace with your data or export directly from your Shiny app.")
    nrun.italic = True

    doc.save(path)

def make_rtf(path):
    text = r"""{\rtf1\ansi
{\b PROPOSED PLAN OF STUDY}\line
Generated: """ + date.today().isoformat() + r"""\line\line
{\b A. Core ({{CORE_TOTAL}} credits)}\line
COURSE\tab TITLE\tab CREDITS\tab GRADE\tab DATE\line
{CODE}\tab {TITLE}\tab {CREDITS}\tab {GRADE}\tab {TERM YEAR}\line\line
{\b B. Additional Required Core ({{ARC_TOTAL}} Hours)}\line
COURSE\tab TITLE\tab CREDITS\tab GRADE\tab DATE\line
{CODE}\tab {TITLE}\tab {CREDITS}\tab {GRADE}\tab {TERM YEAR}\line\line
{\b C. Electives ({{ELECTIVE_TOTAL}} Hours)}\line
COURSE\tab TITLE\tab CREDITS\tab GRADE\tab DATE\line
{CODE}\tab {TITLE}\tab {CREDITS}\tab {GRADE}\tab {TERM YEAR}\line\line
{\b D. Dissertation Hours ({{DISS_TOTAL}})}\line
COURSE\tab TITLE\tab CREDITS\tab GRADE\tab DATE\line
{CODE}\tab {TITLE}\tab {CREDITS}\tab {GRADE}\tab {TERM YEAR}\line\line
{\b SUMMARY OF PROPOSED COURSE OF STUDY CREDIT HOURS}\line
Core\tab {{CORE_TOTAL}}\line
Concentration Core\tab {{ARC_TOTAL}}\line
Electives\tab {{ELECTIVE_TOTAL}}\line
Dissertation\tab {{DISS_TOTAL}}\line
Total\tab {{OVERALL_TOTAL}}\line\line
{\b PROJECTED TIMETABLE}\line
Date Admitted:\tab {{TERM YEAR}}\line
Hours completed to date\tab {{HOURS_COMPLETED}}\line
Completion of Coursework\tab {{TERM YEAR}}\line
Comprehensive exam\tab {{TERM YEAR}}\line
Dissertation proposal\tab {{TERM YEAR}}\line
Semesters of Residency\tab {{RESIDENCY_SEMESTERS}}\line
Completion of dissertation\tab {{TERM YEAR}}\line
Notes\tab {{NOTES}}\line
}"""
    path.write_text(text, encoding="utf-8")

created = None
try:
    import docx  # noqa: F401
    make_docx(docx_path)
    created = ("docx", str(docx_path))
except Exception as e:
    make_rtf(rtf_path)
    created = ("rtf", str(rtf_path), str(e))

created
